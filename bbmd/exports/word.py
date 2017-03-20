import copy
import io
import os

from docx import Document
from docx.shared import Inches

from .. import mpl
from ..bmr.base import BMRBase


class WordReportFactory(object):

    def __init__(self, session, file_):
        f, should_close = self.get_file_object(file_)
        template = os.path.join(os.path.dirname(__file__), 'template.docx')
        doc = Document(template)
        self.create_report(session, doc)
        doc.save(f)
        f.seek(0)
        if should_close:
            f.close()

    def get_file_object(self, file_):
        should_close = False
        if isinstance(file_, basestring):
            file_ = file_.replace('~', os.path.expanduser('~'))
            file_ = os.path.abspath(file_)
            file_ = open(file_, 'w')
            should_close = True
        if not hasattr(file_, 'write'):
            raise ValueError('Must be file-object or valid filename: %s' % file_)
        return file_, should_close

    def create_report(self, session, doc):

        # write heading
        doc.add_heading(session.name, 0)

        # add timestamp
        now = session._get_timestamp().strftime('%b %d, %Y at %I:%M %p')
        txt = 'Report created on {}.'.format(now)
        doc.add_paragraph(txt)

        # add model version
        txt = 'Pystan model version {}.'.format(session.get_pystan_version())
        doc.add_paragraph(txt)

        # write dataset table
        doc.add_heading('Dataset', 1)
        self.dataset_table(doc, session)

        if hasattr(session, 'trend_p_value'):
            self.add_field(
                doc, 'Trend-test p-value',
                self.stringify(session.trend_p_value))

        if hasattr(session, 'trend_z_test'):
            self.add_field(
                doc, 'Z-score',
                self.stringify(session.trend_z_test))

        # write mcmc
        doc.add_heading('BMD Markov chain Monte Carlo settings', 1)
        self.mcmc_settings(doc, session)
        doc.add_page_break()

        # write model fits
        doc.add_heading('Model results', 1)
        for model in session.models:
            self.model_fit(doc, session, model)
            doc.add_page_break()

        # write bmds
        doc.add_heading('BMD results', 1)
        for bmr in BMRBase.get_related_models(session.bmrs):
            self.bmd_estimate(doc, session, bmr)
            doc.add_page_break()

    def dataset_table(self, doc, session):
        doses = session.dataset['d']
        if session.dataset_type in session.DICHOTOMOUS_TYPES:
            n = session.dataset['n']
            incs = session.dataset['y']
            tbl = doc.add_table(rows=n.size + 1, cols=3, style='TableGridA')

            # header
            tbl.rows[0].cells[0].text = 'Dose'
            tbl.rows[0].cells[1].text = 'N'
            tbl.rows[0].cells[2].text = 'Incidence'

            # data rows
            for i, d in enumerate(doses):
                cells = tbl.rows[i + 1].cells
                cells[0].text = str(doses[i])
                cells[1].text = str(n[i])
                cells[2].text = str(incs[i])
        else:
            if session.dataset_type == 'C':
                resps = session.dataset['resp']
                n = session.dataset['n']
                stdevs = session.dataset['stdev']
                tbl = doc.add_table(rows=n.size + 1, cols=4, style='TableGridA')

                # header
                tbl.rows[0].cells[0].text = 'Dose'
                tbl.rows[0].cells[1].text = 'N'
                tbl.rows[0].cells[2].text = 'Response'
                tbl.rows[0].cells[3].text = 'Standard deviation'

                # data rows
                for i, d in enumerate(doses):
                    cells = tbl.rows[i + 1].cells
                    cells[0].text = str(doses[i])
                    cells[1].text = str(n[i])
                    cells[2].text = str(resps[i])
                    cells[3].text = str(stdevs[i])
            else:
                resps = session.dataset['y']
                tbl = doc.add_table(rows=doses.size + 1, cols=2, style='TableGridA')

                # header
                tbl.rows[0].cells[0].text = 'Dose'
                tbl.rows[0].cells[1].text = 'Response'

                # data rows
                for i, d in enumerate(doses):
                    cells = tbl.rows[i + 1].cells
                    cells[0].text = str(doses[i])
                    cells[1].text = str(resps[i])

    def mcmc_settings(self, doc, session):
        self.add_field(
            doc, 'Iterations', format(session.mcmc_iterations, ',d'))
        self.add_field(
            doc, 'Number of chains', str(session.mcmc_num_chains))
        self.add_field(
            doc, 'Warmup fraction', str(session.mcmc_warmup_fraction))
        self.add_field(
            doc, 'Seed', format(session.seed, ',d'))

    def model_fit(self, doc, session, model):
        doc.add_heading(model.name, 2)

        img = self.get_model_equation_image(model)
        doc.add_picture(img, height=Inches(0.35))

        if 'pwr_lbound' in model.data:
            self.add_field(
                doc,
                'Power parameter lower-bound',
                self.stringify(model.data['pwr_lbound']),
            )

        if hasattr(model, 'fit_summary'):
            doc.add_heading('Model fit summary', 3)
            doc.add_paragraph(model.fit_summary, style='code')

            params = (model.plotting.T).tolist()
            drs = self.get_dr_plotting_data(session)
            fig = mpl.model_fit(drs, params)
            doc.add_picture(self.mpl_to_file(fig), width=Inches(5.5))

        if model.predicted_pvalue:
            self.add_field(
                doc,
                'Posterior predictive p-value for model fit',
                self.stringify(model.predicted_pvalue)
            )

        if model.model_weight_scaler:
            self.add_field(
                doc,
                'Model weight',
                '{:.1f}%'.format(model.model_weight_scaler*100.)
            )

        if model.PARAMETERS:
            doc.add_heading('Correlation matrix', 3)
            matrix = self.get_model_correlation_matrix(
                model.PARAMETERS, model.parameter_correlation)
            ncols = len(matrix[0])
            tbl = doc.add_table(rows=ncols, cols=ncols, style='TableGridB')
            for i in range(ncols):
                for j in range(ncols):
                    tbl.rows[i].cells[j].text = matrix[i][j]

            doc.add_heading('Parameter charts', 3)
            fig = mpl.parameter_plots(model.PARAMETERS, model.parameters)
            doc.add_picture(self.mpl_to_file(fig), width=Inches(4.35))

    def get_dr_plotting_data(self, session):
        ds = session.dataset
        doses = ds['d']
        err_low = None
        err_high = None
        if session.dataset_type in session.DICHOTOMOUS_TYPES:
            ct = ds['y'] / ds['n'].astype('float64')
        else:
            if session.dataset_type == 'C':
                ct = ds['resp']
                stdevs = ds['stdev']
                # to do - use more complex CI method
                err_low = err_high = 1.96 * stdevs
            else:
                ct = ds['y']

        return {
            'dose': doses,
            'ct': ct,
            'errLow': err_low,
            'errHigh': err_high,
        }

    def get_model_correlation_matrix(self, parameter_names, parameter_correlation):
        names = list(copy.deepcopy(parameter_names))
        corrs = copy.deepcopy(parameter_correlation)

        for i, corr in enumerate(corrs):
            corrs[i] = [self.stringify(c) if c != 1. else '-' for c in corr]

        for i, name in enumerate(names):
            corrs[i].insert(0, name)

        names.insert(0, '')

        m = [names]
        for corr in corrs:
            m.append(corr)

        return m

    @staticmethod
    def mpl_to_file(plt):
        # convert matplotlib plot into a file-like in-memory object
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        return buf

    @staticmethod
    def get_model_equation_image(model):
        buf = io.BytesIO()
        mpl.render_latex(buf, model.LATEX_EQUATION)
        buf.seek(0)
        return buf

    @staticmethod
    def stringify(val):
        # convert float value into a presentable number
        if isinstance(val, float):
            if val >= 0.001:
                return '{:.3f}'.format(val)
            else:
                return '{:1.3g}'.format(val)
        return str(val)

    @staticmethod
    def add_field(doc, field, content):
        p = doc.add_paragraph()
        r = p.add_run(field + u': ')
        r.bold = True
        p.add_run(content)

    def bmd_estimate(self, doc, session, bmds):

        is_dual = len(bmds) == 2

        doc.add_heading(bmds[0].name, 2)

        doc.add_heading('BMD estimates', 3)

        data = []
        for bmd in bmds:
            kernels = []
            for i, res in enumerate(bmd.results):
                kernels.append(dict(
                    model_name=session.models[i].name,
                    kernel=res['kernel'],
                ))
            data.append(dict(
                dual_type=getattr(bmd, 'DUAL_TYPE', None),
                model_average=bmd.model_average['kernel'],
                models=kernels,
            ))
        fig = mpl.bmd_kernels(is_dual, data)
        width = 4.5 if is_dual else 2.25
        doc.add_picture(self.mpl_to_file(fig), width=Inches(width))

        if is_dual:
            doc.add_heading('BMD summary tables: %s' % bmds[0].DUAL_TYPE, 3)
            self.render_stats_table(doc, session, bmds[0])
            doc.add_heading('BMD summary tables: %s' % bmds[1].DUAL_TYPE, 3)
            self.render_stats_table(doc, session, bmds[1])
        else:
            doc.add_heading('BMD summary tables', 3)
            self.render_stats_table(doc, session, bmds[0])

    def render_stats_table(self, doc, session, bmd):
        data = [d['stats'] for d in bmd.results]
        data.insert(0, bmd.model_average['stats'])

        nrows = 9
        ncols = len(session.models) + 2
        tbl = doc.add_table(rows=nrows, cols=ncols, style='TableGridC')
        content = self.get_stats_table_content(session, bmd)
        for i in range(nrows):
            for j in range(ncols):
                tbl.rows[i].cells[j].text = content[i][j]

    def get_stats_table_content(self, session, bmd):
        # return a 2D list of lists for rendering in table
        content = []

        # insert model columns
        stats = [d['stats'] for d in bmd.results]
        content.append([
            c.name
            for c in session.models])
        content.append([
            self.stringify(p)
            for p in bmd._priors])
        content.append([
            self.stringify(wt)
            for wt in bmd.model_posterior_weights])
        content.append([
            self.stringify(d['p50'])
            for d in stats])
        content.append([
            self.stringify(d['p5'])
            for d in stats])
        content.append([
            self.stringify(d['p25'])
            for d in stats])
        content.append([
            '{} ({})'.format(self.stringify(d['mean']), self.stringify(d['std']))
            for d in stats])
        content.append([
            self.stringify(d['p75'])
            for d in stats])
        content.append([
            self.stringify(d['p95'])
            for d in stats])

        # insert header column
        cells = [
            'Statistic', 'Prior model weight', 'Posterior model weight',
            'BMD (median)', 'BMDL (5%)', '25%',
            'Mean (SD)', '75%', '95%',
        ]
        for i, c in enumerate(cells):
            content[i].insert(0, c)

        # insert model-average column
        stats = bmd.model_average['stats']
        cells = [
            'Model average', 'N/A', 'N/A',
            self.stringify(stats['p50']),
            self.stringify(stats['p5']),
            self.stringify(stats['p25']),
            '{} ({})'.format(self.stringify(
                stats['mean']), self.stringify(stats['std'])),
            self.stringify(stats['p75']),
            self.stringify(stats['p95']),
        ]
        for i, c in enumerate(cells):
            content[i].insert(1, c)

        return content
